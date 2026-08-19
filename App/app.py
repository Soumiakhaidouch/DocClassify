from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, Response
from functools import wraps
from dotenv import load_dotenv
import secrets
import os
import json
import torch
from transformers import XLMRobertaTokenizer, XLMRobertaForSequenceClassification
from pypdf import PdfReader
from docx import Document as DocxDocument
from preprocessor import get_clean_text, detect_language, chunk_text
from auth import auth, close_db, get_db

# ── Limites upload ────────────────────────────────────────────
MAX_FILE_SIZE_MB    = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
ALLOWED_EXTENSIONS  = {".pdf", ".txt", ".docx"}

# ── App ───────────────────────────────────────────────────────
app = Flask(__name__)
# En prod, SECRET_KEY doit être fixe (sinon les sessions sautent à chaque redémarrage) :
# défini dans .env. Le token aléatoire n'est qu'un filet de sécurité pour le dev local.
app.secret_key = os.environ.get("SECRET_KEY") or secrets.token_hex(32)

app.register_blueprint(auth)
app.teardown_appcontext(close_db)

# ── Model ─────────────────────────────────────────────────────

load_dotenv()

MODEL_ID = os.getenv("HF_MODEL_ID")
#HF_TOKEN = os.getenv("HF_TOKEN")

tokenizer = XLMRobertaTokenizer.from_pretrained(
    MODEL_ID,

)

model = XLMRobertaForSequenceClassification.from_pretrained(
    MODEL_ID,
)
model.eval()

CATEGORIES = ['culture', 'finance', 'medical', 'politics', 'sports', 'tech']

# Traduction des catégories selon la langue détectée
CATEGORY_LABELS = {
    'culture':  {'fr': 'Culture',     'ar': 'ثقافة',      'en': 'Culture'},
    'finance':  {'fr': 'Finance',     'ar': 'مالية',      'en': 'Finance'},
    'medical':  {'fr': 'Médical',     'ar': 'طبي',        'en': 'Medical'},
    'politics': {'fr': 'Politique',   'ar': 'سياسة',      'en': 'Politics'},
    'sports':   {'fr': 'Sport',       'ar': 'رياضة',      'en': 'Sports'},
    'tech':     {'fr': 'Technologie', 'ar': 'تكنولوجيا',  'en': 'Technology'},
}

LANG_LABELS = {
    'fr': 'Français',
    'ar': 'العربية',
    'en': 'English',
    'unknown': 'Inconnue',
}

# ── Helpers ───────────────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            flash("Veuillez vous connecter.", "warning")
            return redirect(url_for("auth.login"))
        return f(*args, **kwargs)
    return decorated

def extraire_texte_pdf(file_stream):
    reader = PdfReader(file_stream)
    texte  = ""
    for page in reader.pages:
        if page.extract_text():
            texte += page.extract_text() + "\n"
    return texte


def extraire_texte_docx(file_stream) -> str:
    """Extrait le texte brut d'un fichier .docx (paragraphes + tableaux)."""
    doc    = DocxDocument(file_stream)
    lignes = []
    for para in doc.paragraphs:
        if para.text.strip():
            lignes.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            lignes.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(lignes)


def get_category_label(slug: str, lang: str) -> str:
    """Retourne le label traduit de la catégorie selon la langue."""
    labels = CATEGORY_LABELS.get(slug, {})
    return labels.get(lang, labels.get('fr', slug.capitalize()))


def classify_with_chunks(texte_clean: str) -> tuple[str, float, dict]:
    """
    Découpe le texte en chunks de ≤400 mots, infère sur chacun,
    puis agrège par vote pondéré (somme des probabilités).

    Retourne (category_slug, confidence, all_scores).
    """
    chunks = chunk_text(texte_clean)

    # Accumulateur de scores (somme des softmax par catégorie)
    score_acc = {cat: 0.0 for cat in CATEGORIES}

    for chunk in chunks:
        if not chunk.strip():
            continue
        inputs = tokenizer(
            chunk,
            return_tensors="pt",
            truncation=True,
            padding=True,
            max_length=512,
        )
        with torch.no_grad():
            outputs = model(**inputs)

        probs = torch.nn.functional.softmax(outputs.logits, dim=1)[0]
        for i, cat in enumerate(CATEGORIES):
            score_acc[cat] += float(probs[i])

    # Normaliser pour obtenir des probabilités moyennes
    n = len(chunks) or 1
    all_scores  = {cat: score_acc[cat] / n for cat in CATEGORIES}
    cat_slug    = max(all_scores, key=all_scores.__getitem__)
    confidence  = all_scores[cat_slug]

    return cat_slug, confidence, all_scores


def store_document_and_classification(
    user_id: int,
    file_name: str,
    file_type: str,
    raw_text: str,
    lang_detected: str,
    category_slug: str,
    confidence: float,
    all_scores: dict,
    file_size: int = 0,
):
    """
    Insère le document uploadé et son résultat de classification dans la BD.
    Retourne l'id du document créé.
    """
    db  = get_db()
    cur = db.cursor()

    char_count  = len(raw_text)
    token_count = len(raw_text.split())

    # 1. Insérer dans documents (le contenu brut est stocké pour l'aperçu dans la modale "Voir")
    cur.execute(
        """
        INSERT INTO documents
            (user_id, file_name, file_type, status, lang_detected, char_count, token_count, file_size, content)
        VALUES (%s, %s, %s, 'done', %s, %s, %s, %s, %s)
        RETURNING id
        """,
        (user_id, file_name, file_type, lang_detected, char_count, token_count, file_size, raw_text),
    )
    doc_id = cur.fetchone()["id"]

    # 2. Récupérer l'id de la catégorie
    cur.execute("SELECT id FROM categories WHERE slug = %s", (category_slug,))
    row = cur.fetchone()
    cat_id = row["id"] if row else None

    # 3. Insérer dans classification
    cur.execute(
        """
        INSERT INTO classification
            (document_id, category_id, confidence, all_scores, model_used)
        VALUES (%s, %s, %s, %s, %s)
        """,
        (doc_id, cat_id, round(confidence, 4), json.dumps(all_scores), "xlm-roberta"),
    )

    # 4. Audit log
    cur.execute(
        """
        INSERT INTO audit_logs (user_id, document_id, action, details)
        VALUES (%s, %s, 'classify', %s)
        """,
        (user_id, doc_id, json.dumps({"category": category_slug, "confidence": round(confidence, 4), "lang": lang_detected})),
    )

    db.commit()
    cur.close()
    return doc_id


# ── Page d'accueil : redirige vers login ou classify ─────────
@app.route("/")
def home():
    if "user_id" in session:
        return redirect(url_for("classify_document"))
    return redirect(url_for("auth.login"))


# ── Classification ────────────────────────────────────────────
@app.route("/classify_document", methods=["GET", "POST"])
def classify_document():
    result_data = None   # dict avec toutes les infos à passer au template

    if request.method == "POST":
        if "user_id" not in session:
            flash("Veuillez vous connecter pour classifier un document.", "warning")
            return redirect(url_for("auth.login", next=url_for("classify_document")))

        texte_brut = request.form.get("texte", "").strip()
        file_name  = "saisie_manuelle"
        file_type  = "text"
        file_size  = 0

        if "file" in request.files and request.files["file"].filename != "":
            file      = request.files["file"]
            file_name = file.filename
            fname_low = file_name.lower()

            # ── Validation extension ──
            ext = os.path.splitext(fname_low)[1]
            if ext not in ALLOWED_EXTENSIONS:
                flash(f"Format non supporté. Formats acceptés : {', '.join(sorted(ALLOWED_EXTENSIONS))}.", "warning")
                return render_template("index.html", result_data=None)

            # ── Validation taille (10 MB max) ──
            file.seek(0, 2)                   # aller à la fin
            file_size = file.tell()
            file.seek(0)                      # revenir au début
            if file_size > MAX_FILE_SIZE_BYTES:
                flash(f"Fichier trop volumineux. Taille maximale autorisée : {MAX_FILE_SIZE_MB} Mo.", "warning")
                return render_template("index.html", result_data=None)

            # ── Extraction du texte ──
            if fname_low.endswith(".pdf"):
                file_type  = "pdf"
                texte_brut = extraire_texte_pdf(file)
            elif fname_low.endswith(".docx"):
                file_type  = "docx"
                texte_brut = extraire_texte_docx(file)
            else:
                file_type  = "txt"
                texte_brut = file.read().decode("utf-8", errors="ignore")

        if not texte_brut:
            flash("Veuillez fournir un texte ou uploader un fichier.", "warning")
            return render_template("index.html", result_data=None)

        # ── Détection de langue ──
        lang_detected = detect_language(texte_brut)

        # ── Nettoyage + inférence par chunks ──
        texte_clean = get_clean_text(texte_brut)
        cat_slug, confidence, all_scores = classify_with_chunks(texte_clean)

        # ── Stockage en BD ──
        doc_id = store_document_and_classification(
            user_id       = session["user_id"],
            file_name     = file_name,
            file_type     = file_type,
            raw_text      = texte_brut,
            lang_detected = lang_detected,
            category_slug = cat_slug,
            confidence    = confidence,
            all_scores    = all_scores,
            file_size     = file_size,
        )

        result_data = {
            "doc_id":       doc_id,
            "category_slug": cat_slug,
            "category_label": get_category_label(cat_slug, lang_detected),
            "confidence":   confidence,
            "lang":         lang_detected,
            "lang_label":   LANG_LABELS.get(lang_detected, lang_detected.upper()),
            "all_scores":   all_scores,
            "category_labels_translated": {
                slug: get_category_label(slug, lang_detected)
                for slug in CATEGORIES
            },
        }

    logged_in = "user_id" in session
    return render_template("index.html", result_data=result_data, logged_in=logged_in)


# ── Historique ────────────────────────────────────────────────
@app.route("/document_history")
@login_required
def document_history():
    db  = get_db()
    cur = db.cursor()
    cur.execute(
        """
        SELECT
            d.id, d.file_name, d.file_type, d.lang_detected,
            d.char_count, d.file_size, d.upload_date,
            cat.slug   AS category_slug,
            cat.name_fr AS category_fr,
            cat.name_ar AS category_ar,
            c.confidence
        FROM documents d
        LEFT JOIN classification c ON c.document_id = d.id
        LEFT JOIN categories cat   ON cat.id = c.category_id
        WHERE d.user_id = %s
        ORDER BY d.upload_date DESC
        LIMIT 50
        """,
        (session["user_id"],),
    )
    docs = cur.fetchall()
    cur.close()
    return render_template("history.html", docs=docs)



# ── Document detail API (pour le modal) ──────────────────────
@app.route("/api/document/<int:doc_id>")
@login_required
def api_get_document(doc_id):
    db  = get_db()
    cur = db.cursor()
    cur.execute(
        """
        SELECT
            d.id, d.file_name, d.file_type, d.lang_detected,
            d.char_count, d.token_count, d.upload_date, d.content,
            cat.slug       AS category_slug,
            cat.name_fr    AS category_fr,
            cat.name_ar    AS category_ar,
            c.confidence,
            c.all_scores,
            c.model_used
        FROM documents d
        LEFT JOIN classification c ON c.document_id = d.id
        LEFT JOIN categories cat   ON cat.id = c.category_id
        WHERE d.id = %s AND d.user_id = %s
        """,
        (doc_id, session["user_id"]),
    )
    doc = cur.fetchone()
    cur.close()

    if not doc:
        return jsonify({"error": "Document introuvable"}), 404

    all_scores = doc["all_scores"]
    if isinstance(all_scores, str):
        try:
            all_scores = json.loads(all_scores)
        except Exception:
            all_scores = {}

    return jsonify({
        "id":            doc["id"],
        "file_name":     doc["file_name"],
        "file_type":     doc["file_type"],
        "lang_detected": doc["lang_detected"],
        "char_count":    doc["char_count"],
        "token_count":   doc["token_count"],
        "content":       doc["content"],
        "upload_date":   doc["upload_date"].strftime('%d/%m/%Y à %H:%M') if doc["upload_date"] else None,
        "category_slug": doc["category_slug"],
        "category_fr":   doc["category_fr"],
        "category_ar":   doc["category_ar"],
        "confidence":    doc["confidence"],
        "all_scores":    all_scores,
        "model_used":    doc["model_used"],
    })



# ── Suppression d'un document ────────────────────────────────
@app.route("/api/document/<int:doc_id>/delete", methods=["DELETE"])
@login_required
def api_delete_document(doc_id):
    db  = get_db()
    cur = db.cursor()
    # Vérifier que le document appartient à l'utilisateur
    cur.execute(
        "SELECT id FROM documents WHERE id = %s AND user_id = %s",
        (doc_id, session["user_id"]),
    )
    if not cur.fetchone():
        cur.close()
        return jsonify({"error": "Document introuvable"}), 404

    # Supprimer classification + audit + document (CASCADE recommandé en BD)
    cur.execute("DELETE FROM classification WHERE document_id = %s", (doc_id,))
    cur.execute("DELETE FROM audit_logs    WHERE document_id = %s", (doc_id,))
    cur.execute("DELETE FROM documents     WHERE id = %s AND user_id = %s",
                (doc_id, session["user_id"]))
    db.commit()
    cur.close()
    return jsonify({"success": True})


@app.route("/browse_categories")
@login_required
def browse_categories():
    """Page d'affichage des catégories sous forme de dossiers avec leurs documents."""
    db  = get_db()
    cur = db.cursor()
 
    # Récupérer tous les documents de l'utilisateur groupés par catégorie
    cur.execute(
        """
        SELECT
            d.id, d.file_name, d.file_type, d.lang_detected,
            d.char_count, d.upload_date,
            cat.slug        AS category_slug,
            cat.name_fr     AS category_fr,
            cat.name_ar     AS category_ar,
            c.confidence
        FROM documents d
        LEFT JOIN classification c   ON c.document_id = d.id
        LEFT JOIN categories cat     ON cat.id = c.category_id
        WHERE d.user_id = %s
        ORDER BY cat.slug, d.upload_date DESC
        """,
        (session["user_id"],),
    )
    rows = cur.fetchall()
    cur.close()
 
    # Grouper par slug de catégorie
    from collections import defaultdict
    docs_by_category = defaultdict(list)
    for row in rows:
        slug = row["category_slug"] or "unknown"
        docs_by_category[slug].append(row)
 
    # Nombre de docs par catégorie (pour les compteurs sur les dossiers)
    counts = {slug: len(docs) for slug, docs in docs_by_category.items()}
 
    return render_template(
        "classe.html",
        docs_by_category=docs_by_category,
        counts=counts,
    )


@app.context_processor
def inject_doc_count():
    """Rend le nombre de documents disponible dans TOUS les templates (sidebar)."""
    if "user_id" not in session:
        return dict(doc_count=0)
    try:
        db = get_db()
        cur = db.cursor()
        cur.execute(
            "SELECT COUNT(*) AS cnt FROM documents WHERE user_id = %s",
            (session["user_id"],)
        )
        count = cur.fetchone()["cnt"]
        cur.close()
    except Exception:
        count = 0
    return dict(doc_count=count)

if __name__ == "__main__":
    app.run(debug=True)